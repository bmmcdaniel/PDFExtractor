"""Text extraction in multiple formats: Markdown, Plain Text, XHTML, DOCX, FVTT."""

import json
import logging
import re
import secrets
import xml.etree.ElementTree as ET
from collections.abc import Callable
from pathlib import Path

import pymupdf
import pymupdf4llm
from docx import Document as DocxDocument

from core.pdf_handler import PDFHandler

log = logging.getLogger(__name__)

# Register XHTML as the default namespace so ET serializes without a prefix.
# This intentionally modifies the process-wide ET namespace registry.
ET.register_namespace('', 'http://www.w3.org/1999/xhtml')
_XHTML_NS = 'http://www.w3.org/1999/xhtml'

_XHTML_CONT_HEADER = (
    '<?xml version="1.0"?>\n'
    '<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.0 Strict//EN"'
    ' "http://www.w3.org/TR/xhtml1/DTD/xhtml1-strict.dtd">\n'
    '<html xmlns="http://www.w3.org/1999/xhtml">\n'
    '<head>\n'
    '<style>\n'
    'p{white-space:pre-wrap}\n'
    '</style>\n'
    '</head>\n'
    '<body>\n'
)
_XHTML_CONT_TRAILER = '</body>\n</html>\n'

_MD_HEADER_RE = re.compile(r'^(#{1,4} +)(.*)', re.MULTILINE)

# XML 1.0 forbids character references to control chars other than tab/LF/CR.
# PyMuPDF emits these (e.g. &#x8;) for non-printable glyphs in PDFs.
_INVALID_XML_CHARREF_RE = re.compile(r'&#x([0-9A-Fa-f]+);|&#([0-9]+);')
# \xad (soft hyphen) is handled separately by _SOFT_HYPHEN_RE below.
_INVALID_XML_RAW_RE = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f￾￿]')

# Soft hyphen (U+00AD) marks an optional line-break position in the PDF.
# PyMuPDF often emits a space after the soft hyphen to represent the line
# break, so we must strip the soft hyphen AND any immediately following
# whitespace to avoid leaving ghost spaces inside words ("wear ing").
_SOFT_HYPHEN_RE = re.compile(r'\xad\s*')

# PyMuPDF emits a whitespace-only bold run at the start of paragraphs that
# follow section headings in many PDFs (e.g. "<p><b> </b>actual text").
# This is a layout artifact — the bold space comes from the PDF's column or
# indent formatting and carries no semantic content.
_LEADING_WS_INLINE_RE = re.compile(r'(<p>)<(?:b|i|strong|em)>\s+</(?:b|i|strong|em)>')


def _strip_soft_hyphens(text: str) -> str:
    return _SOFT_HYPHEN_RE.sub('', text)


# Matches the first non-whitespace character at the start of the string or
# immediately following any whitespace.  Used by _title_case() to capitalise
# only true word-initial letters, avoiding the str.title() behaviour of also
# capitalising after apostrophes ("Boy'S") and hyphens ("Jack-In-The-Box").
_WORD_INITIAL_RE = re.compile(r'(^|\s)(\S)')


def _title_case(s: str) -> str:
    """Title-case *s* using whitespace-only word boundaries.

    Differs from str.title() in that apostrophes and hyphens are NOT treated
    as word separators, so possessives and contractions capitalise correctly:
      "THE WITCH'S HOVEL"  → "The Witch's Hovel"   (not "The Witch'S Hovel")
      "JACK-IN-THE-BOX"   → "Jack-in-the-box"       (not "Jack-In-The-Box")
    """
    return _WORD_INITIAL_RE.sub(
        lambda m: m.group(1) + m.group(2).upper(), s.lower()
    )


def _title_case_element(el) -> None:
    """Apply _title_case() to all text nodes within el, preserving child markup."""
    if el.text:
        el.text = _title_case(el.text)
    for child in el:
        _title_case_element(child)
        if child.tail:
            child.tail = _title_case(child.tail)


def _normalize_xhtml_headers(content: str) -> str:
    """Apply _title_case() to h1–h4 text content in an XHTML document string."""
    html_idx = content.find('<html')
    if html_idx == -1:
        return content
    prolog = content[:html_idx]
    root = ET.fromstring(content[html_idx:])
    for level in range(1, 5):
        for el in root.iter(f'{{{_XHTML_NS}}}h{level}'):
            _title_case_element(el)
    return prolog + ET.tostring(root, encoding='unicode')


def _normalize_markdown_headers(text: str) -> str:
    """Apply _title_case() to h1–h4 header lines in Markdown."""
    return _MD_HEADER_RE.sub(lambda m: m.group(1) + _title_case(m.group(2)), text)


def _sanitize_xml(text: str) -> str:
    """Strip character references and raw bytes invalid in XML 1.0."""
    text = _SOFT_HYPHEN_RE.sub('', text)
    text = _LEADING_WS_INLINE_RE.sub(r'\1', text)
    def _keep_ref(m: re.Match) -> str:
        cp = int(m.group(1), 16) if m.group(1) else int(m.group(2))
        if (cp in (0x9, 0xA, 0xD)
                or 0x20 <= cp <= 0xD7FF
                or 0xE000 <= cp <= 0xFFFD
                or 0x10000 <= cp <= 0x10FFFF):
            return m.group(0)
        return ''
    text = _INVALID_XML_CHARREF_RE.sub(_keep_ref, text)
    return _INVALID_XML_RAW_RE.sub('', text)


def _strip_page_div(xhtml: str) -> str:
    """Return inner content of the outer <div> wrapper from a page XHTML fragment."""
    open_end = xhtml.find('>')
    if open_end == -1:
        return xhtml
    close_start = xhtml.rfind('</div>')
    if close_start == -1:
        return xhtml
    return xhtml[open_end + 1:close_start]


_PAGE_DIV_RE = re.compile(r'<div\s+id=["\']page0["\']>', re.IGNORECASE)
_XMLNS_ATTR_RE = re.compile(r'\s+xmlns(?::\w+)?="[^"]*"')


def _extract_page_div_content(full_xhtml: str) -> str:
    """Extract the inner HTML of <div id="page0"> from a full XHTML document.

    Used for FVTT output where the XHTML wrapper (prolog, html/head/body tags)
    must be stripped and only the page content kept.  Namespace declarations
    are removed so the result is plain HTML suitable for Foundry.
    """
    m = _PAGE_DIV_RE.search(full_xhtml)
    if not m:
        return full_xhtml
    start = m.end()
    end = full_xhtml.rfind('</div>')
    inner = full_xhtml[start:end] if end > start else full_xhtml[start:]
    return _XMLNS_ATTR_RE.sub('', inner)


def _add_formatted_runs(para, text: str) -> None:
    """Parse Markdown inline formatting into python-docx Runs.

    Handles ***bold italic***, **bold**, and *italic*.
    """
    pattern = re.compile(
        r"\*\*\*(.+?)\*\*\*"
        r"|\*\*(.+?)\*\*"
        r"|\*(.+?)\*"
        r"|([^*]+|\*+)"
    )
    for match in pattern.finditer(text):
        bold_italic, bold, italic, plain = match.groups()
        if bold_italic:
            run = para.add_run(bold_italic)
            run.bold = True
            run.italic = True
        elif bold:
            run = para.add_run(bold)
            run.bold = True
        elif italic:
            run = para.add_run(italic)
            run.italic = True
        elif plain:
            para.add_run(plain)


class TextExtractor:
    """Handles text extraction in multiple formats."""

    def __init__(self, pdf_handler: PDFHandler):
        """Initialize with an open PDFHandler."""
        self._handler = pdf_handler

    def extract_markdown(
        self,
        pages: list[int],
        output_path: str,
        progress_callback: Callable[[int, int], None] | None = None,
        per_page: bool = False,
        normalize_headers: bool = False,
        filename_map: dict[int, str] | None = None,
    ) -> None:
        """Extract text as Markdown using pymupdf4llm."""
        if per_page:
            base = Path(output_path)
            for i, page_num in enumerate(pages):
                md_text = pymupdf4llm.to_markdown(
                    self._handler.document,
                    pages=[page_num],
                    page_chunks=False,
                    show_progress=False,
                    table_strategy="lines_strict",
                )
                md_text = _strip_soft_hyphens(md_text)
                if normalize_headers:
                    md_text = _normalize_markdown_headers(md_text)
                hex_stem = (filename_map or {}).get(page_num)
                stem = f"{base.stem}_{hex_stem}" if hex_stem else f"{base.stem}_{i + 1:03d}"
                out_path = base.parent / f"{stem}{base.suffix}"
                out_path.write_text(md_text, encoding="utf-8")
                if progress_callback:
                    progress_callback(i + 1, len(pages))
        else:
            if progress_callback:
                progress_callback(0, len(pages))
            md_text = pymupdf4llm.to_markdown(
                self._handler.document,
                pages=pages,
                page_chunks=False,
                show_progress=False,
                table_strategy="lines_strict",
            )
            md_text = _strip_soft_hyphens(md_text)
            if normalize_headers:
                md_text = _normalize_markdown_headers(md_text)
            Path(output_path).write_text(md_text, encoding="utf-8")
            if progress_callback:
                progress_callback(len(pages), len(pages))

    def extract_plain_text(
        self,
        pages: list[int],
        output_path: str,
        progress_callback: Callable[[int, int], None] | None = None,
        per_page: bool = False,
        filename_map: dict[int, str] | None = None,
    ) -> None:
        """Extract text as plain text."""
        doc = self._handler.document
        if per_page:
            base = Path(output_path)
            for i, page_num in enumerate(pages):
                text = _strip_soft_hyphens(doc[page_num].get_text("text"))
                hex_stem = (filename_map or {}).get(page_num)
                stem = f"{base.stem}_{hex_stem}" if hex_stem else f"{base.stem}_{i + 1:03d}"
                out_path = base.parent / f"{stem}{base.suffix}"
                out_path.write_text(text, encoding="utf-8")
                if progress_callback:
                    progress_callback(i + 1, len(pages))
        else:
            text_parts = []
            for i, page_num in enumerate(pages):
                text_parts.append(_strip_soft_hyphens(doc[page_num].get_text("text")))
                if progress_callback:
                    progress_callback(i + 1, len(pages))
            Path(output_path).write_text("\f".join(text_parts), encoding="utf-8")

    def extract_xhtml(
        self,
        pages: list[int],
        output_path: str,
        progress_callback: Callable[[int, int], None] | None = None,
        include_images: bool = False,
        per_page: bool = False,
        normalize_headers: bool = False,
        filename_map: dict[int, str] | None = None,
        page_xhtml_transform: Callable[[str], str] | None = None,
    ) -> None:
        """Extract text as XHTML."""
        self._extract_xhtml_impl(
            pages, output_path,
            include_images=include_images,
            per_page=per_page,
            normalize_headers=normalize_headers,
            progress_callback=progress_callback,
            filename_map=filename_map,
            page_xhtml_transform=page_xhtml_transform,
        )

    def _extract_xhtml_impl(
        self,
        pages: list[int],
        output_path: str,
        include_images: bool,
        per_page: bool = False,
        normalize_headers: bool = False,
        progress_callback: Callable[[int, int], None] | None = None,
        filename_map: dict[int, str] | None = None,
        page_xhtml_transform: Callable[[str], str] | None = None,
    ) -> None:
        doc = self._handler.document
        flags = pymupdf.TEXT_PRESERVE_LIGATURES | pymupdf.TEXT_PRESERVE_WHITESPACE
        if include_images:
            flags |= pymupdf.TEXT_PRESERVE_IMAGES

        base = Path(output_path)

        if per_page:
            header = pymupdf.ConversionHeader("xhtml", filename=self._handler.filename).lstrip("\n")
            trailer = pymupdf.ConversionTrailer("xhtml")
            for i, page_num in enumerate(pages):
                page_xhtml = _sanitize_xml(doc[page_num].get_text("xhtml", flags=flags))
                content = "\n".join([header, page_xhtml, trailer])
                if normalize_headers:
                    content = _normalize_xhtml_headers(content)
                hex_stem = (filename_map or {}).get(page_num)
                if hex_stem and page_xhtml_transform is not None:
                    content = page_xhtml_transform(content)
                stem = f"{base.stem}_{hex_stem}" if hex_stem else f"{base.stem}_{i + 1:03d}"
                out_path = base.parent / f"{stem}{base.suffix}"
                out_path.write_text(content, encoding="utf-8")
                if progress_callback:
                    progress_callback(i + 1, len(pages))
        else:
            body_parts = []
            for i, page_num in enumerate(pages):
                page_xhtml = _sanitize_xml(doc[page_num].get_text("xhtml", flags=flags))
                body_parts.append(_strip_page_div(page_xhtml))
                if progress_callback:
                    progress_callback(i + 1, len(pages))
            content = _XHTML_CONT_HEADER + "".join(body_parts) + _XHTML_CONT_TRAILER
            if normalize_headers:
                content = _normalize_xhtml_headers(content)
            base.write_text(content, encoding="utf-8")

    def extract_docx(
        self,
        pages: list[int],
        output_path: str,
        progress_callback: Callable[[int, int], None] | None = None,
        normalize_headers: bool = False,
    ) -> None:
        """Extract text as a Word document using python-docx."""
        if progress_callback:
            progress_callback(0, len(pages))
        chunks = pymupdf4llm.to_markdown(
            self._handler.document,
            pages=pages,
            page_chunks=True,
            table_strategy="lines_strict",
        )

        docx_doc = DocxDocument()

        for i, chunk in enumerate(chunks):
            md_text = _strip_soft_hyphens(chunk["text"])

            for line in md_text.split("\n"):
                stripped = line.strip()
                if not stripped:
                    continue

                if stripped.startswith("#### "):
                    text = stripped[5:]
                    docx_doc.add_heading(_title_case(text) if normalize_headers else text, level=4)
                elif stripped.startswith("### "):
                    text = stripped[4:]
                    docx_doc.add_heading(_title_case(text) if normalize_headers else text, level=3)
                elif stripped.startswith("## "):
                    text = stripped[3:]
                    docx_doc.add_heading(_title_case(text) if normalize_headers else text, level=2)
                elif stripped.startswith("# "):
                    text = stripped[2:]
                    docx_doc.add_heading(_title_case(text) if normalize_headers else text, level=1)
                else:
                    para = docx_doc.add_paragraph()
                    _add_formatted_runs(para, stripped)

            if i < len(chunks) - 1:
                docx_doc.add_page_break()

            if progress_callback:
                progress_callback(i + 1, len(chunks))

        docx_doc.save(output_path)

    def extract_fvtt(
        self,
        pages: list[int],
        output_path: str,
        progress_callback: Callable[[int, int], None] | None = None,
        profile=None,                        # ExtractionProfile | None
        filename_map: dict[int, str] | None = None,
        page_xhtml_transform: Callable[[str], str] | None = None,
        normalize_headers: bool = False,
        include_images: bool = False,
        stem_to_title: Callable[[str], str] | None = None,
        stem_to_semantic_keys: Callable[[str], dict] | None = None,
    ) -> None:
        """Extract text as a Foundry VTT journal JSON file.

        One Foundry JournalEntry is written containing one JournalEntryPage
        per journal group defined by *profile*.  If *profile* is None, each
        PDF page becomes its own journal page (equivalent to toc_level=0).

        *filename_map*          page index → stem; drives titling and semantic
                                key extraction for "page" strategy sections.
        *page_xhtml_transform*  called on each page whose section has
                                transform=True (e.g. DCB hex restructuring).
        *stem_to_title*         converts a stem to a display title.
        *stem_to_semantic_keys* extracts link keys from a stem.
        """
        from core.extraction_profile import (
            ExtractionProfile, SectionRule, JournalGroup,
            build_journal_groups, resolve_links,
        )

        doc = self._handler.document

        if profile is None:
            profile = ExtractionProfile(
                journal_name=Path(output_path).stem,
                sections=[SectionRule("Content", 0, doc.page_count - 1, "page")],
            )

        groups: list[JournalGroup] = build_journal_groups(
            profile, doc, pages, filename_map, stem_to_title, stem_to_semantic_keys,
        )

        journal_id = secrets.token_hex(8)

        # page_id_map: 0-based PDF page → journal page _id
        page_id_map: dict[int, str] = {}
        for grp in groups:
            for pg in grp.pages:
                page_id_map[pg] = grp.page_id

        # semantic_index: key_type → {key_value → page_id}
        semantic_index: dict[str, dict[str, str]] = {}
        for grp in groups:
            for key_type, key_value in grp.semantic_keys.items():
                semantic_index.setdefault(key_type, {})[key_value] = grp.page_id

        flags = pymupdf.TEXT_PRESERVE_LIGATURES | pymupdf.TEXT_PRESERVE_WHITESPACE
        if include_images:
            flags |= pymupdf.TEXT_PRESERVE_IMAGES
        xhtml_header = pymupdf.ConversionHeader(
            "xhtml", filename=self._handler.filename
        ).lstrip("\n")
        xhtml_trailer = pymupdf.ConversionTrailer("xhtml")

        total = len(groups)
        fvtt_pages = []

        for gi, grp in enumerate(groups):
            parts: list[str] = []
            for pg in grp.pages:
                raw = _sanitize_xml(doc[pg].get_text("xhtml", flags=flags))
                full_xhtml = "\n".join([xhtml_header, raw, xhtml_trailer])
                if normalize_headers:
                    full_xhtml = _normalize_xhtml_headers(full_xhtml)
                if grp.transform and page_xhtml_transform is not None:
                    full_xhtml = page_xhtml_transform(full_xhtml)
                parts.append(_extract_page_div_content(full_xhtml))

            content = "".join(parts)
            content = _strip_soft_hyphens(content)

            if profile.link_rules:
                content = resolve_links(
                    content, profile.link_rules,
                    journal_id, page_id_map, semantic_index,
                )

            fvtt_pages.append({
                "_id": grp.page_id,
                "name": grp.title,
                "type": "text",
                "sort": (gi + 1) * 100000,
                "title": {"show": True, "level": 1},
                "text": {
                    # format 1 = JOURNAL_ENTRY_PAGE_FORMATS.HTML in Foundry v10+
                    "content": content,
                    "format": 1,
                },
                "image": {},
                "video": {"controls": True, "volume": 0.5},
                "src": None,
                "system": {},
                "ownership": {"default": -1},
                "flags": {},
            })

            if progress_callback:
                progress_callback(gi + 1, total)

        journal = {
            "_id": journal_id,
            "name": profile.journal_name,
            "ownership": {"default": 0},
            "flags": {},
            "pages": fvtt_pages,
            "_stats": {
                "coreVersion": "12.0.0",
                "systemVersion": "0",
                "createdTime": 0,
                "modifiedTime": 0,
                "lastModifiedBy": "",
            },
        }

        Path(output_path).write_text(
            json.dumps(journal, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )
